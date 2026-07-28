"""Compact Word body-block reads and guarded localized patch operations."""
from __future__ import annotations

import hashlib
import zipfile

from core import store

import office_tx

DEFAULT_LIMIT = 50
MAX_LIMIT = 500
MAX_PATCH_OPS = 100
SNIPPET_CHARS = 160


class WordError(Exception):
    pass


def _docx():
    try:
        import docx
        return docx
    except ImportError as exc:
        raise WordError(".docx support needs the 'python-docx' package") from exc


def _document_preflight(path: str) -> None:
    try:
        with zipfile.ZipFile(path) as package:
            xml = package.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile) as exc:
        raise WordError("file is not a valid Word document") from exc
    for marker, label in (
        (b"<w:ins", "tracked revisions"),
        (b"<w:del", "tracked revisions"),
        (b"<w:sdt", "content controls"),
        (b"<w:altChunk", "altChunk content"),
    ):
        if marker in xml:
            raise WordError(f"Word mutation does not safely support {label}")


def _kind(paragraph) -> str:
    style = getattr(paragraph.style, "name", "") or ""
    if style.startswith("Heading"):
        suffix = style[len("Heading"):].strip()
        return "h" + (suffix or "?")
    if style.startswith("List"):
        return "li"
    return "p"


def _block_id(index: int, paragraph) -> str:
    payload = "\x1f".join((_kind(paragraph), paragraph.style.name or "", paragraph.text))
    digest = hashlib.sha256(payload.encode("utf-8", "replace")).hexdigest()[:8]
    return f"p{index}-{digest}"


def _blocks(document):
    return [
        {
            "id": _block_id(index, paragraph),
            "index": index,
            "kind": _kind(paragraph),
            "style": paragraph.style.name or "",
            "text": paragraph.text,
            "paragraph": paragraph,
        }
        for index, paragraph in enumerate(document.paragraphs, 1)
    ]


def outline(path: str, *, offset: int = 0, limit: int = DEFAULT_LIMIT) -> dict:
    if offset < 0 or limit < 1 or limit > MAX_LIMIT:
        raise WordError(f"offset must be >= 0 and limit must be 1..{MAX_LIMIT}")
    office_tx._package_preflight(path, mutating=False)
    document = _docx().Document(path)
    blocks = _blocks(document)
    page = blocks[offset:offset + limit]
    return {
        "hash": store.file_sha256(path),
        "blocks": [
            [item["id"], item["kind"],
             item["text"][:SNIPPET_CHARS] + (
                 "..." if len(item["text"]) > SNIPPET_CHARS else ""
             )]
            for item in page
        ],
        "offset": offset,
        "returned": len(page),
        "more": offset + len(page) < len(blocks),
        "unsupported": {
            "tables": len(document.tables),
            "headers_footers_and_nested_parts": "not enumerated",
        },
    }


def read_blocks(path: str, ids: list[str]) -> dict:
    if not ids or len(ids) > MAX_LIMIT:
        raise WordError(f"request 1..{MAX_LIMIT} block IDs")
    office_tx._package_preflight(path, mutating=False)
    document = _docx().Document(path)
    mapping = {item["id"]: item for item in _blocks(document)}
    missing = [value for value in ids if value not in mapping]
    if missing:
        raise WordError(f"stale or unknown Word block ID(s): {', '.join(missing[:5])}")
    return {
        "hash": store.file_sha256(path),
        "blocks": [
            {
                "id": value,
                "kind": mapping[value]["kind"],
                "style": mapping[value]["style"],
                "text": mapping[value]["text"],
            }
            for value in ids
        ],
    }


def _paragraph_is_complex(paragraph) -> bool:
    xml = paragraph._p.xml
    return any(token in xml for token in (
        "<w:fldChar", "<w:drawing", "<w:object", "<w:bookmarkStart",
        "<w:commentRangeStart", "<w:hyperlink",
    ))


def _insert_paragraph(target, *, before: bool, style: str, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    element = OxmlElement("w:p")
    if before:
        target._p.addprevious(element)
    else:
        target._p.addnext(element)
    paragraph = Paragraph(element, target._parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def patch(
    path: str,
    operations: list,
    *,
    expected_sha256: str,
    dry_run: bool = False,
) -> dict:
    if not expected_sha256:
        raise WordError("Word patch needs --expected-file-hash from outline/read-blocks")
    if not isinstance(operations, list) or not 1 <= len(operations) <= MAX_PATCH_OPS:
        raise WordError(f"Word patch needs 1..{MAX_PATCH_OPS} operations")
    resolved = {}

    def build_plan(live_path):
        _document_preflight(live_path)
        document = _docx().Document(live_path)
        mapping = {item["id"]: item for item in _blocks(document)}
        seen = set()
        allowed = {
            "replace_text": {"op", "id", "find", "replace"},
            "replace_block": {"op", "id", "text", "style"},
            "insert_before": {"op", "id", "blocks"},
            "insert_after": {"op", "id", "blocks"},
            "append": {"op", "blocks"},
            "delete_block": {"op", "id"},
        }
        for number, operation in enumerate(operations, 1):
            if not isinstance(operation, dict):
                raise WordError(f"Word patch operation {number} must be an object")
            op = operation.get("op")
            if op not in allowed:
                raise WordError(f"unsupported Word patch operation: {op!r}")
            unknown = set(operation) - allowed[op]
            if unknown:
                raise WordError(f"unknown fields for {op}: {', '.join(sorted(unknown))}")
            block_id = operation.get("id")
            if block_id:
                if block_id not in mapping:
                    raise WordError(f"stale or unknown Word block ID: {block_id}")
                if block_id in seen:
                    raise WordError("a Word block may be targeted only once per patch")
                seen.add(block_id)
                if op in ("replace_text", "replace_block", "delete_block") \
                        and _paragraph_is_complex(mapping[block_id]["paragraph"]):
                    raise WordError("target paragraph contains unsupported complex content")
            if op == "replace_text":
                find = operation.get("find")
                if not isinstance(find, str) or not find:
                    raise WordError("replace_text needs non-empty find text")
                if mapping[block_id]["text"].count(find) != 1:
                    raise WordError("replace_text must match exactly once in its block")
            if op in ("replace_block",):
                if not isinstance(operation.get("text"), str):
                    raise WordError("replace_block needs text")
            if op in ("insert_before", "insert_after", "append"):
                blocks = operation.get("blocks")
                if not isinstance(blocks, list) or not blocks:
                    raise WordError(f"{op} needs a non-empty blocks array")
                for block in blocks:
                    if not (isinstance(block, list) and len(block) == 2
                            and all(isinstance(value, str) for value in block)):
                        raise WordError("inserted blocks must be [style, text] pairs")
        resolved["count"] = len(operations)
        return office_tx.MutationPlan(
            "patch",
            {"operations": len(operations)},
            {"affected": len(operations)},
        )

    def apply(stage, _plan):
        document = _docx().Document(stage)
        mapping = {item["id"]: item["paragraph"] for item in _blocks(document)}
        for operation in operations:
            op = operation["op"]
            paragraph = mapping.get(operation.get("id"))
            if op == "replace_text":
                full = "".join(run.text for run in paragraph.runs)
                updated = full.replace(operation["find"], operation.get("replace", ""), 1)
                remaining = updated
                runs = list(paragraph.runs)
                for index, run in enumerate(runs):
                    if index == len(runs) - 1:
                        run.text = remaining
                    else:
                        width = len(run.text)
                        run.text, remaining = remaining[:width], remaining[width:]
            elif op == "replace_block":
                style = operation.get("style") or paragraph.style.name
                paragraph.text = operation["text"]
                paragraph.style = style
            elif op in ("insert_before", "insert_after"):
                before = op == "insert_before"
                anchor = paragraph
                sequence = operation["blocks"] if before else list(reversed(operation["blocks"]))
                for style, text in sequence:
                    _insert_paragraph(anchor, before=before, style=style, text=text)
            elif op == "append":
                for style, text in operation["blocks"]:
                    document.add_paragraph(text, style=style or None)
            elif op == "delete_block":
                element = paragraph._element
                element.getparent().remove(element)
        document.save(stage)

    def validate(stage, _plan):
        _document_preflight(stage)
        _docx().Document(stage)
        return {"patched": resolved["count"]}

    try:
        return office_tx.execute_mutation(
            path, operation="patch", plan=build_plan, apply=apply,
            validate=validate, expected_sha256=expected_sha256,
            dry_run=dry_run,
        )
    except office_tx.TransactionError as exc:
        raise WordError(str(exc)) from exc
